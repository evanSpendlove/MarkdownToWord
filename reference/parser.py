# Structure
# Styling rules object
# Parser for markdown
# Main script

# TODO(evanSpendlove): Add parsing for italics
# TODO(evanSpendlove): Add parsing for hyperlinks

from docx import Document
from styles import *

class MarkdownParser:
    def __init__(self, doc = Document()):
        self.document = doc
        self.style = Style()

    def applyStyles(self, tabs, runs):
        para = self.document.add_paragraph('')
        para = self.style.setIndent(para, tabs)
        for r in runs:
            style, content = r[0], r[1]
            run = para.add_run(content)
            self.style.styleRun(run, style)
        return para

    def parseTabs(self, line):
        tabs = 0
        tabMarker = line[0] if line[0] in ' \t' else '\0'
        while tabs < len(line) and line[tabs] == tabMarker:
            tabs += 1
        if tabMarker == ' ':
            tabs //= 4
        return tabs

    def parseWords(self, line, tabs):
        runs = []
        start, end = 0, 0
        bold = False
        while end < len(line):
            if line[end] == '*':
                if line[end - 1] == '*':
                    if not bold:
                        bold = True
                        runs.append(("NONE", line[start:end-1]))
                    else:
                        bold = False
                        runs.append(("BOLD", line[start:end-1]))
                    start = end + 1
            end += 1
        runs.append(("NONE", line[start:end]))
        print(runs)
        return self.applyStyles(tabs, runs)

    def parseHeader(self, line, tabs):
        level = 0
        while line[level] == '#': level += 1
        para = self.parseWords(line[level+1:], tabs)
        return self.style.setHeader(para, level)

    def parseList(self, line, tabs):
        line = line[line.index('-')+2:]
        para = self.parseWords(line, tabs)
        return self.style.setList(para, tabs)

    def parseLine(self, line):
        if line == '':
            return
        tabs = self.parseTabs(line)
        if '# ' in line:
            return self.parseHeader(line, tabs)
        if '- ' in line:
            return self.parseList(line, tabs)
        return self.parseWords(line, tabs)
