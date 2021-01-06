import docx

class Paragraph_List(object):
    def __init__(self,*args): #args are: doc,item1,ordered,style,fmt (style and fmt optional)
        self.args = args
        self.doc =self.args[0]
        self.item1 = self.args[1]
        self.ordered = self.args[2]
        self.place = {}
        List_dict = {'Roman':[['I','II','III','IV','V','VI','VII','VIII','IIX','IX','X'],['A','B','C','D','E','F',
            'G','H','I','J','K','L'],['1','2','3','4','5','6','7','8','9','10'],['a','b','c','d','e','f','g','h',
            'i'],['i','ii','iii','iv','v','vi','vii','viii','iix','ix','x']],'ABC':[['A','B','C','D','E','F','G',
            'H','I','J','K','L'],['1','2','3','4','5','6','7','8','9','10'],['a','b','c','d','e','f','g','h','i'],
            ['i','ii','iii','iv','v','vi','vii','viii','iix','ix','x']],'123':[['1','2','3','4','5','6','7','8','9','10'
            ],['a','b','c','d','e','f','g','h','i'],['i','ii','iii','iv','v','vi','vii','viii','iix','ix','x']],'Bullet':[
            ['●','○','•','◦']]}
        if self.ordered == True:
            if len(self.args) < 4:
                self.fmt = List_dict['Roman']
            elif self.args[3] == 'Custom':
                self.fmt = self.args[4]
            else:
                self.fmt = List_dict[self.args[3]]
            self.p = self.doc.add_paragraph(self.fmt[0][0]+'. ' + self.item1 + '\n')
        else:
            self.fmt = List_dict['Bullet']
            self.p = self.doc.add_paragraph(self.fmt[0][0]+ ' ' + self.item1 + '\n')
        # self.doc.add_paragraph(item1, style=self.level)
  
    def add_item(self, item, level):
        self.level = level
        self.place[1] = 0
        if self.level ==1:
            sp = ""
        else:
            sp = "    "
            sp = sp *(self.level -1)
        if self.level in self.place:
            self.place[self.level] += 1
        else:
            self.place[self.level] = 0
        if self.ordered ==True:
            self.p.add_run(sp + self.fmt[self.level - 1][self.place[self.level]] + '. ' + item + '\n')
        else:
            self.p.add_run(sp + self.fmt[0][self.level - 1]+ ' ' + item + '\n')

document = docx.Document()
mylist = Paragraph_List(document, 'item 1', True)
mylist.add_item('Level 2 Item 1',2)
mylist.add_item('Level 2 Item 2',2)
mylist.add_item('Level 2 Item 3',2)
mylist.add_item('Level 3 Item 1',3)
mylist.add_item('Level 3 Item 2',3)
mylist.add_item('Level 3 Item 3',3)
mylist.add_item('Level 3 Item 4',3)
mylist.add_item('Level 3 Item 5',3)
mylist.add_item('Level 1 Item 2',1)

mylist2 = Paragraph_List(document, 'Bullet Level 1 Item 1', False)
mylist2.add_item('Level 1 Item2',1)
mylist2.add_item('Level 2 Item1',2)
mylist2.add_item('Level 1 Item3',1)
mylist2.add_item('Level 2 Item1',2)
mylist2.add_item('Level 2 Item2',2)
mylist2.add_item('Level 2 Item3',2)
mylist2.add_item('Level 3 Item1',3)
mylist2.add_item('Level 4 Item1',4)
mylist2.add_item('Level 4 Item2',4)
mylist2.add_item('Level 2 Item4',2)

document.save('new.docx') 
