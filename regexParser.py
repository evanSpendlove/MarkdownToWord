import re
from docx import Document
from styles import *

class MarkdownParser:
    def __init__(self, doc = Document()):
        self.doc = doc
        self.style = Style()
        return

    def parseLine(self, line):
        if self.parseHeader(line) is None:
            self.parseContent(line)

    def parseContent(self, content):
       para =  self.doc.add_paragraph('')
       runs = self.parseBold(content)
       if runs is None:
           para.add_run(content)
           return para
       for r in runs:
           para.add_run(r[0])
           para.add_run(r[1]).bold = True
           para.add_run(r[2])
       return para

    def parseBold(self, content):
        bold = r"([^*]*)\*\*(?P<bold>[^*]+)\*\*([^*]*)"
        matches = re.findall(bold, content)
        return matches if len(matches) > 0 else None

    def parseHeader(self, line):
        header = r"^(?P<header>#+)(?: +)(?P<content>.*)$"
        match = re.match(header, line)
        if match is None:
            return None
        height = len(match.group('header'))
        content = self.parseContent(match.group('content'))
        return self.style.header(content, height)

md = MarkdownParser()
tests = [
        "**John**",
        "Hi",
        "Hello, my **name** is...",
        "Hello, my **name** is **John**"
        ]
for t in tests:
    match = md.parseContent(t)
    if match is not None: 
        print(t)
        for m in match.runs:
            if m.bold:
                print("BOLD" + m.text)
            else:
                print(m.text)
