from docx import *
from docx.enum.text import WD_ALIGN_PARAGRAPH

class Styler:
    def __init__(self):
        self.doc = Document()

    def paragraph(self, para):
        if not para.header:
            return self.doc.add_paragraph('')
        else:
            return self.header('', para.level)

    def applyStyle(self, para, run):
        r = para.add_run(run.content)
        if run.style == 'BOLD':
            r.bold = True
        elif run.style == 'ITALIC':
            r.italic = True
        return r

    def header(self, content, level):
        # If Heading 1, align center
        para = self.doc.add_paragraph(content)
        if level > 9:
            raise Exception('Heading level cannot be greater than 9')
        para.style = f"Heading {level}"
        if level == 1:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return para

    def setIndent(self, para, tabs):
        TAB_INDENT = 0.63       # Cm
        para.paragraph_format.left_indent = Cm(TAB_INDENT * tabs)
        return para

    def setList(self, para, tabs):
        listStyle = 'List Bullet'
        # if tabs > 0: listStyle += f" {tabs + 1}"
        para.style = listStyle
        return para

    def styleRun(self, run, style):
        if style == "BOLD":
            run.bold = True
        return run
